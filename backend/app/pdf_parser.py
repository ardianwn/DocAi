"""
PDF Parser Module

This module handles PDF document parsing and text extraction for the DocAI system.
"""

import os
import logging
from typing import List, Dict, Any, Optional
from PyPDF2 import PdfReader
from docx import Document
import aiofiles

logger = logging.getLogger(__name__)


class PDFParser:
    """
    A class to handle parsing of PDF and other document formats.
    
    This class provides methods to extract text content from various document
    formats including PDF, DOCX, and plain text files.
    """
    
    def __init__(self):
        """Initialize the PDF Parser."""
        self.supported_formats = ['.pdf', '.docx', '.txt']
        logger.info("PDFParser initialized")
    
    async def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a PDF file and extract text content.
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            Dict[str, Any]: Dictionary containing extracted text and metadata
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            Exception: If parsing fails
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            logger.info(f"Parsing PDF file: {file_path}")
            
            # Extract text from PDF
            text_content = []
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                metadata['num_pages'] = len(pdf_reader.pages)
                metadata['title'] = pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else ''
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append({
                            'page': page_num + 1,
                            'content': page_text.strip()
                        })
            
            result = {
                'text_content': text_content,
                'metadata': metadata,
                'file_path': file_path,
                'format': 'pdf'
            }
            
            logger.info(f"Successfully parsed PDF: {len(text_content)} pages extracted")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            raise
    
    async def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a DOCX file and extract text content.
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            Dict[str, Any]: Dictionary containing extracted text and metadata
        """
        try:
            logger.info(f"Parsing DOCX file: {file_path}")
            
            doc = Document(file_path)
            text_content = []
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content.append({
                        'paragraph': para_num + 1,
                        'content': paragraph.text.strip()
                    })
            
            metadata = {
                'num_paragraphs': len(text_content),
                'total_paragraphs': len(doc.paragraphs)
            }
            
            result = {
                'text_content': text_content,
                'metadata': metadata,
                'file_path': file_path,
                'format': 'docx'
            }
            
            logger.info(f"Successfully parsed DOCX: {len(text_content)} paragraphs extracted")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
    
    async def parse_text_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a plain text file.
        
        Args:
            file_path (str): Path to the text file
            
        Returns:
            Dict[str, Any]: Dictionary containing extracted text and metadata
        """
        try:
            logger.info(f"Parsing text file: {file_path}")
            
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            
            lines = content.split('\n')
            text_content = []
            
            for line_num, line in enumerate(lines):
                if line.strip():
                    text_content.append({
                        'line': line_num + 1,
                        'content': line.strip()
                    })
            
            metadata = {
                'num_lines': len(text_content),
                'total_lines': len(lines),
                'file_size': os.path.getsize(file_path)
            }
            
            result = {
                'text_content': text_content,
                'metadata': metadata,
                'file_path': file_path,
                'format': 'txt'
            }
            
            logger.info(f"Successfully parsed text file: {len(text_content)} lines extracted")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {str(e)}")
            raise
    
    async def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a document based on its file extension.
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            Dict[str, Any]: Dictionary containing extracted text and metadata
            
        Raises:
            ValueError: If file format is not supported
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            if file_extension == '.pdf':
                return await self.parse_pdf(file_path)
            elif file_extension == '.docx':
                return await self.parse_docx(file_path)
            elif file_extension == '.txt':
                return await self.parse_text_file(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            raise
    
    def get_supported_formats(self) -> List[str]:
        """
        Get list of supported document formats.
        
        Returns:
            List[str]: List of supported file extensions
        """
        return self.supported_formats.copy()


# Utility function for backwards compatibility
async def parse_pdf_file(file_path: str) -> Dict[str, Any]:
    """
    Utility function to parse a PDF file.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        Dict[str, Any]: Parsed content and metadata
    """
    parser = PDFParser()
    return await parser.parse_pdf(file_path)